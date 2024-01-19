# Copyright (c) 2024 Emina Mahmutbegovic
#
# All rights reserved.
# This software is the proprietary information of Emina Mahmutbegovic
# Unauthorized sharing of this file is strictly prohibited
import os

import numpy as np
import datetime
import ast
import csv
from docx import Document
from PyQt5.QtWidgets import QFileDialog


def generate_mse_report(scores):
    scores = np.asarray(scores)

    # Create a new Word document
    doc = Document()

    # Get the current date and time
    current_datetime = datetime.datetime.now()

    # Format the datetime as a string
    timestamp_string = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")

    heading = 'Report_' + timestamp_string

    doc_name = heading + '.docx'

    # Create the full output file path
    output_filename = os.path.join(get_file_path('output'), doc_name)

    # Add a title to the document
    doc.add_heading(heading, level=1)

    # Add the MSE section to the document
    doc.add_heading('MSE:', level=2)
    doc.add_paragraph(f'Scores Mean: {scores.mean():.4f} A² +- {2 * scores.std():.4f} A²')
    doc.add_paragraph(f'Scores Min: {scores.min():.4f} A², Scores Max: {scores.max():.4f} A²')

    # Save the document to the specified output filename
    doc.save(output_filename)


# Generic function for generating report for a string array
def generate_report(string_input, prefix=''):
    # Create a new Word document
    doc = Document()

    # Get the current date and time
    current_datetime = datetime.datetime.now()

    # Format the datetime as a string
    timestamp_string = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")

    heading = prefix + 'Report_' + timestamp_string

    doc_name = heading + '.docx'

    # Create the full output file path
    output_filename = os.path.join(get_file_path('output'), doc_name)

    # Add a title to the document
    doc.add_heading(heading, level=1)

    # Add the MSE section to the document
    doc.add_heading(prefix, level=2)
    doc.add_paragraph(string_input)

    # Save the document to the specified output filename
    doc.save(output_filename)


def get_file_path(relative_path):
    # Get the current working directory
    current_directory = os.getcwd()

    # Create the absolute path by joining the current directory and the relative path
    return os.path.join(current_directory, relative_path)


def open_save_file_dialog(window):
    # Open a file dialog to choose where to save the file
    options = QFileDialog.Options()
    filename, _ = QFileDialog.getSaveFileName(window, "Save File", "",
                                              "All Files (*);;Text Files (*.txt)",
                                              options=options)

    return filename


# Upload csv style function
def upload_csv(parent):
    options = QFileDialog.Options()
    options |= QFileDialog.ReadOnly
    file_name, _ = QFileDialog.getOpenFileName(
        parent,
        "Open CSV File",
        "",
        "CSV Files (*.csv);;All Files (*)",
        options=options
    )

    if file_name:
        print(f"Selected CSV File: {file_name}")
        return file_name


# Save text file
def save_file(window, string_content):
    filename = open_save_file_dialog(window)

    if filename:
        # Write content to the file
        with open(filename, 'w') as file:
            file.write(string_content)
        print(f"File saved: {filename}")


# Save csv file for history reports
def save_history_reports_csv_file(window, string_content):
    filename = open_save_file_dialog(window)

    # Parse the string into actual Python dictionaries using ast.literal_eval
    parsed_data = [ast.literal_eval(item) for item in string_content]

    if filename:
        # Create a CSV file
        with open(f"{filename}_training_report.csv", 'w', newline='') as csvfile:
            writer = csv.writer(csvfile)
            # Write the header
            writer.writerow(['Loss', 'Accuracy'])

            # Write the data
            for entry in parsed_data:
                loss = entry['loss'][0]
                accuracy = entry['accuracy'][0]
                writer.writerow([loss, accuracy])


# Save csv file for evaluation reports
def save_evaluation_reports_csv_file(window, string_content):
    filename = open_save_file_dialog(window)

    if filename:
        # Create a CSV file
        with open(f"{filename}_evaluation_report.csv", 'w', newline='') as csvfile:
            writer = csv.writer(csvfile)
            # Write the header
            writer.writerow(['Loss', 'Accuracy'])

            # Write each data row
            for row in string_content:
                writer.writerow(row)
