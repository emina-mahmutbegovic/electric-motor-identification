# Copyright (c) 2024 Emina Mahmutbegovic
#
# All rights reserved.
# This software is the proprietary information of Emina Mahmutbegovic
# Unauthorized sharing of this file is strictly prohibited
import os

import numpy as np
import datetime
from docx import Document


def generate_mse_report(scores):
    scores = np.asarray(scores)

    # Create a new Word document
    doc = Document()

    # Get the current date and time
    current_datetime = datetime.datetime.now()

    # Format the datetime as a string
    timestamp_string = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")

    heading = 'Report_' + timestamp_string

    doc_name = heading + '.docx'

    # Create the full output file path
    output_filename = os.path.join(get_file_path('output'), doc_name)

    # Add a title to the document
    doc.add_heading(heading, level=1)

    # Add the MSE section to the document
    doc.add_heading('MSE:', level=2)
    doc.add_paragraph(f'Scores Mean: {scores.mean():.4f} A² +- {2 * scores.std():.4f} A²')
    doc.add_paragraph(f'Scores Min: {scores.min():.4f} A², Scores Max: {scores.max():.4f} A²')

    # Save the document to the specified output filename
    doc.save(output_filename)


# Generic function for generating report for a string array
def generate_report(string_input, prefix=''):
    # Create a new Word document
    doc = Document()

    # Get the current date and time
    current_datetime = datetime.datetime.now()

    # Format the datetime as a string
    timestamp_string = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")

    heading = prefix + 'Report_' + timestamp_string

    doc_name = heading + '.docx'

    # Create the full output file path
    output_filename = os.path.join(get_file_path('output'), doc_name)

    # Add a title to the document
    doc.add_heading(heading, level=1)

    # Add the MSE section to the document
    doc.add_heading(prefix, level=2)
    doc.add_paragraph(string_input)

    # Save the document to the specified output filename
    doc.save(output_filename)


def get_file_path(relative_path):
    # Get the current working directory
    current_directory = os.getcwd()

    # Create the absolute path by joining the current directory and the relative path
    return os.path.join(current_directory, relative_path)
